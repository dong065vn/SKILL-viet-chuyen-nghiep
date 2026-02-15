#!/usr/bin/env python3
"""
Export Data - Xuất dữ liệu crawl ra Excel, CSV, Word.
Tự động tạo folder output riêng cho mỗi lần xuất.

Usage:
    python export_data.py --preview <data_file.json>
    python export_data.py --export <format> --input <data_file.json> [--output-dir <folder>]

Formats: excel, csv, word, all
"""

import argparse
import csv
import json
import os
import re
import sys
import time
from urllib.parse import urlparse

try:
    from openpyxl import Workbook
    from openpyxl.styles import (
        Alignment,
        Border,
        Font,
        PatternFill,
        Side,
    )
    from openpyxl.utils import get_column_letter
except ImportError:
    print("❌ Thiếu openpyxl. Chạy: pip install openpyxl")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Thiếu python-docx. Chạy: pip install python-docx")


# ─── Output Directory ────────────────────────────────────────────────────────


def create_output_dir(data, custom_dir=None):
    """Create a dedicated output directory for this crawl session."""
    if custom_dir:
        output_dir = custom_dir
    else:
        # Generate folder name from source URL + timestamp
        source_url = data.get("source_url", "unknown")
        parsed = urlparse(source_url)
        domain = parsed.netloc.replace("www.", "").replace(".", "_")
        # Clean domain name
        domain = re.sub(r'[^\w\-]', '_', domain)
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        
        output_dir = os.path.join("output", f"{domain}_{timestamp}")
    
    os.makedirs(output_dir, exist_ok=True)
    print(f"  📁 Thư mục output: {os.path.abspath(output_dir)}")
    return output_dir


# ─── Load Data ────────────────────────────────────────────────────────────────


def load_data(data_file):
    """Load crawled data from JSON file."""
    if not os.path.exists(data_file):
        print(f"❌ Không tìm thấy file: {data_file}")
        sys.exit(1)

    with open(data_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    return data


# ─── Preview ──────────────────────────────────────────────────────────────────


def preview_data(data_file, limit=10):
    """Preview crawled data in a formatted table."""
    data = load_data(data_file)

    print(f"\n{'═' * 60}")
    print(f"👀 PREVIEW DỮ LIỆU")
    print(f"{'═' * 60}")
    print(f"  Nguồn: {data.get('source_url', 'N/A')}")
    print(f"  Ngày crawl: {data.get('crawled_at', 'N/A')}")
    print(f"  Tổng trang: {data.get('total_pages', 0)}")
    print(f"  Tổng records: {data.get('total_records', 0)}")
    print(f"  Trường: {', '.join(data.get('fields', []))}")
    print(f"{'═' * 60}\n")

    records = data.get("data", [])
    if not records:
        print("  ⚠️ Không có dữ liệu để hiện.")
        return

    fields = data.get("fields", list(records[0].keys()))
    show_records = records[:limit]

    # Calculate column widths
    col_widths = {}
    for field in fields:
        col_widths[field] = max(
            len(field),
            max((len(str(r.get(field, ""))[:40]) for r in show_records), default=0),
        )
        col_widths[field] = min(col_widths[field], 40)

    # Print header
    header = " │ ".join(f.ljust(col_widths[f])[:col_widths[f]] for f in fields)
    separator = "─┼─".join("─" * col_widths[f] for f in fields)
    print(f"  {header}")
    print(f"  {separator}")

    # Print rows
    for i, record in enumerate(show_records, 1):
        row = " │ ".join(
            str(record.get(f, ""))[:col_widths[f]].ljust(col_widths[f])
            for f in fields
        )
        print(f"  {row}")

    if len(records) > limit:
        print(f"\n  ... và {len(records) - limit} records nữa")

    print(f"\n{'═' * 60}")
    print(f"  Chọn định dạng xuất:")
    print(f"  [1] Excel (.xlsx) - Có format bảng, màu sắc")
    print(f"  [2] CSV (.csv)    - Nhẹ, mở được mọi nơi")
    print(f"  [3] Word (.docx)  - Có format bảng đẹp")
    print(f"  [4] Tất cả        - Xuất cả 3 file")
    print(f"{'═' * 60}")


# ─── Export Excel ─────────────────────────────────────────────────────────────


def export_excel(data_file, output_dir):
    """Export data to Excel with professional formatting."""
    data = load_data(data_file)
    records = data.get("data", [])
    fields = data.get("fields", list(records[0].keys()) if records else [])

    if not records:
        print("❌ Không có dữ liệu để xuất.")
        return None

    output_path = os.path.join(output_dir, f"crawl_data.xlsx")

    wb = Workbook()
    ws = wb.active
    ws.title = "Dữ liệu Crawl"

    # ── Title row ──
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(fields))
    title_cell = ws.cell(row=1, column=1)
    title_cell.value = f"Dữ liệu từ: {data.get('source_url', 'N/A')}"
    title_cell.font = Font(name="Arial", size=14, bold=True, color="1F4E79")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 35

    # ── Info row ──
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(fields))
    info_cell = ws.cell(row=2, column=1)
    info_cell.value = f"Ngày crawl: {data.get('crawled_at', 'N/A')} | Tổng: {len(records)} records | Trang: {data.get('total_pages', 0)}"
    info_cell.font = Font(name="Arial", size=10, italic=True, color="666666")
    info_cell.alignment = Alignment(horizontal="center")
    ws.row_dimensions[2].height = 25

    # ── Header row ──
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="D9E2F3"),
        right=Side(style="thin", color="D9E2F3"),
        top=Side(style="thin", color="D9E2F3"),
        bottom=Side(style="thin", color="D9E2F3"),
    )

    for col, field in enumerate(fields, 1):
        cell = ws.cell(row=4, column=col, value=field)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    ws.row_dimensions[4].height = 30

    # ── Data rows with zebra stripes ──
    even_fill = PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid")
    data_font = Font(name="Arial", size=10)
    data_alignment = Alignment(vertical="top", wrap_text=True)

    for row_idx, record in enumerate(records, 5):
        for col_idx, field in enumerate(fields, 1):
            value = str(record.get(field, ""))
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = data_font
            cell.alignment = data_alignment
            cell.border = thin_border
            if (row_idx - 5) % 2 == 1:
                cell.fill = even_fill

    # ── Auto-width columns ──
    for col_idx, field in enumerate(fields, 1):
        max_length = len(field)
        for row_idx in range(5, min(len(records) + 5, 55)):
            cell_value = str(ws.cell(row=row_idx, column=col_idx).value or "")
            max_length = max(max_length, min(len(cell_value), 50))
        ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 4

    # ── Freeze panes ──
    ws.freeze_panes = "A5"

    # ── Auto filter ──
    ws.auto_filter.ref = f"A4:{get_column_letter(len(fields))}{len(records) + 4}"

    wb.save(output_path)
    file_size = os.path.getsize(output_path) / 1024

    print(f"  ✅ Excel: {output_path} ({file_size:.1f} KB)")
    return output_path


# ─── Export CSV ───────────────────────────────────────────────────────────────


def export_csv(data_file, output_dir):
    """Export data to CSV with UTF-8 BOM encoding."""
    data = load_data(data_file)
    records = data.get("data", [])
    fields = data.get("fields", list(records[0].keys()) if records else [])

    if not records:
        print("❌ Không có dữ liệu để xuất.")
        return None

    output_path = os.path.join(output_dir, f"crawl_data.csv")

    # UTF-8 BOM for proper Vietnamese display in Excel
    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for record in records:
            writer.writerow(record)

    file_size = os.path.getsize(output_path) / 1024
    print(f"  ✅ CSV: {output_path} ({file_size:.1f} KB)")
    return output_path


# ─── Export Word ──────────────────────────────────────────────────────────────


def export_word(data_file, output_dir):
    """Export data to Word document with professional table formatting."""
    data = load_data(data_file)
    records = data.get("data", [])
    fields = data.get("fields", list(records[0].keys()) if records else [])

    if not records:
        print("❌ Không có dữ liệu để xuất.")
        return None

    output_path = os.path.join(output_dir, f"crawl_data.docx")

    doc = Document()

    # ── Document styles ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Title ──
    title = doc.add_heading("Báo cáo Dữ liệu Crawl", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Info paragraph ──
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        f"Nguồn: {data.get('source_url', 'N/A')}\n"
        f"Ngày crawl: {data.get('crawled_at', 'N/A')}\n"
        f"Tổng records: {len(records)} | Tổng trang: {data.get('total_pages', 0)}"
    )
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")  # Spacer

    # ── Data table ──
    table = doc.add_table(rows=1, cols=len(fields))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, field in enumerate(fields):
        cell = header_row.cells[i]
        cell.text = field
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Header background color
        cell_element = cell._element
        tc_pr = cell_element.get_or_add_tcPr()
        from docx.oxml.ns import qn
        shading = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): "1F4E79",
            qn("w:val"): "clear",
        })
        tc_pr.append(shading)

    # Data rows
    for row_idx, record in enumerate(records):
        row = table.add_row()
        for col_idx, field in enumerate(fields):
            value = str(record.get(field, ""))
            cell = row.cells[col_idx]
            cell.text = value[:200]  # Limit long text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            # Zebra stripe
            if row_idx % 2 == 1:
                cell_element = cell._element
                tc_pr = cell_element.get_or_add_tcPr()
                shading = tc_pr.makeelement(qn("w:shd"), {
                    qn("w:fill"): "F2F7FB",
                    qn("w:val"): "clear",
                })
                tc_pr.append(shading)

    # ── Footer ──
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(f"Xuất bởi Crawl Engine • {time.strftime('%d/%m/%Y %H:%M')}")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    file_size = os.path.getsize(output_path) / 1024

    print(f"  ✅ Word: {output_path} ({file_size:.1f} KB)")
    return output_path


# ─── CLI ──────────────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(
        description="📄 Export Data - Xuất dữ liệu crawl ra Excel/CSV/Word",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--preview", metavar="FILE", help="Preview dữ liệu từ file JSON")
    group.add_argument("--export", choices=["excel", "csv", "word", "all"], help="Xuất dữ liệu")

    parser.add_argument("--input", metavar="FILE", help="File dữ liệu JSON đầu vào")
    parser.add_argument("--output-dir", metavar="DIR", help="Thư mục output (mặc định: tự tạo)")

    args = parser.parse_args()

    if args.preview:
        preview_data(args.preview)

    elif args.export:
        if not args.input:
            print("❌ Cần chỉ định --input khi export")
            sys.exit(1)

        data_file = args.input
        data = load_data(data_file)

        # Create output directory
        output_dir = create_output_dir(data, args.output_dir)

        # Copy JSON data to output dir too
        json_dest = os.path.join(output_dir, "crawl_data.json")
        with open(json_dest, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"\n📄 XUẤT DỮ LIỆU")
        print(f"{'═' * 50}")

        if args.export == "excel" or args.export == "all":
            export_excel(data_file, output_dir)

        if args.export == "csv" or args.export == "all":
            export_csv(data_file, output_dir)

        if args.export == "word" or args.export == "all":
            export_word(data_file, output_dir)

        print(f"{'═' * 50}")
        print(f"📁 Tất cả file lưu tại: {os.path.abspath(output_dir)}")
        print(f"✅ Hoàn tất!")


if __name__ == "__main__":
    main()
