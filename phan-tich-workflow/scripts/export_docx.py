#!/usr/bin/env python3
"""
export_docx.py — Xuất tài liệu DOCX từ kiến thức Tam Giác Thành Công

Khi đã thu thập đủ kiến thức từ 3 skill (Tư Duy Đúng, Kiến Thức Đúng, Công Cụ Đúng),
script này render template DOCX chuyên nghiệp.

Usage:
    python export_docx.py --input knowledge.json --output output.docx
    python export_docx.py --sample  # Generate sample DOCX with test data
"""

import argparse
import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("❌ python-docx chưa được cài đặt.")
    print("   Chạy: pip install python-docx")
    sys.exit(1)


# ─── Brand Colors ───────────────────────────────────────────────
COLORS = {
    "primary": RGBColor(0x6C, 0x63, 0xFF),      # Purple - main brand
    "tu_duy": RGBColor(0xEF, 0x44, 0x44),        # Red - Tư Duy Đúng
    "kien_thuc": RGBColor(0x3B, 0x82, 0xF6),     # Blue - Kiến Thức Đúng
    "cong_cu": RGBColor(0x22, 0xC5, 0x5E),       # Green - Công Cụ Đúng
    "dark": RGBColor(0x1F, 0x29, 0x37),           # Dark text
    "gray": RGBColor(0x6B, 0x72, 0x80),           # Gray text
    "white": RGBColor(0xFF, 0xFF, 0xFF),          # White
}


def setup_styles(doc):
    """Thiết lập styles cho document."""
    # Title style
    style = doc.styles['Title']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(28)
    font.color.rgb = COLORS["primary"]
    font.bold = True

    # Heading 1
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(20)
    font.color.rgb = COLORS["dark"]
    font.bold = True

    # Heading 2
    style = doc.styles['Heading 2']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.color.rgb = COLORS["primary"]
    font.bold = True

    # Normal text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = COLORS["dark"]


def add_cover_page(doc, topic, date_str):
    """Tạo trang bìa."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🏆 TAM GIÁC THÀNH CÔNG")
    run.font.size = Pt(32)
    run.font.color.rgb = COLORS["primary"]
    run.font.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phân Tích & Lộ Trình Đạt Mục Tiêu")
    run.font.size = Pt(16)
    run.font.color.rgb = COLORS["gray"]

    doc.add_paragraph()

    # Topic
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Chủ đề: {topic}")
    run.font.size = Pt(20)
    run.font.color.rgb = COLORS["dark"]
    run.font.bold = True

    doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Ngày tạo: {date_str}")
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS["gray"]

    # Framework summary
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🧠 Tư Duy Đúng 70%  |  📘 Kiến Thức Đúng 10%  |  🔧 Công Cụ Đúng 20%")
    run.font.size = Pt(11)
    run.font.color.rgb = COLORS["gray"]

    # Page break
    doc.add_page_break()


def add_toc_placeholder(doc):
    """Thêm mục lục placeholder."""
    doc.add_heading("Mục Lục", level=1)
    p = doc.add_paragraph()
    run = p.add_run("[Mục lục sẽ được tự động tạo khi mở file trong Word]")
    run.font.color.rgb = COLORS["gray"]
    run.font.italic = True
    p = doc.add_paragraph()
    run = p.add_run("Tip: Trong Word, bấm Ctrl+A → F9 để cập nhật mục lục.")
    run.font.color.rgb = COLORS["gray"]
    run.font.size = Pt(9)
    doc.add_page_break()


def add_styled_table(doc, headers, rows, color=None):
    """Thêm bảng có style."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = COLORS["white"]

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # Spacing after table


def add_section_tu_duy(doc, data):
    """Phần 1: Tư Duy Đúng (70%)."""
    doc.add_heading("Phần 1: 🧠 Tư Duy Đúng (70%) — Thuận Nguyên Lý", level=1)

    # Nguyên lý
    if data.get("nguyen_ly"):
        doc.add_heading("Nguyên lý nền tảng", level=2)
        for item in data["nguyen_ly"]:
            doc.add_paragraph(str(item), style='List Bullet')

    # First Principles
    if data.get("first_principles"):
        doc.add_heading("First Principles Analysis", level=2)
        for item in data["first_principles"]:
            doc.add_paragraph(str(item), style='List Bullet')

    # Mental Models
    if data.get("mental_models"):
        doc.add_heading("Mental Models áp dụng", level=2)
        for item in data["mental_models"]:
            doc.add_paragraph(str(item), style='List Bullet')

    # 5 Whys
    if data.get("five_whys"):
        doc.add_heading("Phân tích 5 Whys", level=2)
        for i, why in enumerate(data["five_whys"], 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Why {i}: ")
            run.font.bold = True
            p.add_run(str(why))

    # Pareto
    if data.get("pareto_top20"):
        doc.add_heading("Pareto 80/20 — Top 20% quan trọng nhất", level=2)
        for item in data["pareto_top20"]:
            doc.add_paragraph(str(item), style='List Bullet')

    # Removed assumptions
    if data.get("removed_assumptions"):
        doc.add_heading("Assumptions đã loại bỏ", level=2)
        for item in data["removed_assumptions"]:
            doc.add_paragraph(f"❌ {item}", style='List Bullet')

    doc.add_page_break()


def add_section_kien_thuc(doc, data):
    """Phần 2: Kiến Thức Đúng (10%)."""
    doc.add_heading("Phần 2: 📘 Kiến Thức Đúng (10%) — Học Nhanh Nhờ AI", level=1)

    # Knowledge gaps
    if data.get("knowledge_gaps"):
        doc.add_heading("Knowledge Gap Analysis", level=2)
        headers = ["Lĩnh vực", "Hiện tại", "Cần đạt", "Gap"]
        rows = []
        for gap in data["knowledge_gaps"]:
            if isinstance(gap, dict):
                rows.append([
                    gap.get("field", ""),
                    str(gap.get("current", "")),
                    str(gap.get("target", "")),
                    str(gap.get("gap", ""))
                ])
            else:
                rows.append([str(gap), "", "", ""])
        if rows:
            add_styled_table(doc, headers, rows)

    # Key concepts
    if data.get("key_concepts"):
        doc.add_heading("Key Concepts", level=2)
        for concept in data["key_concepts"]:
            if isinstance(concept, dict):
                p = doc.add_paragraph()
                run = p.add_run(f"• {concept.get('name', '')}")
                run.font.bold = True
                if concept.get("description"):
                    p.add_run(f" — {concept['description']}")
                if concept.get("level"):
                    run2 = p.add_run(f" [{concept['level']}]")
                    run2.font.color.rgb = COLORS["gray"]
            else:
                doc.add_paragraph(str(concept), style='List Bullet')

    # Curated resources
    if data.get("curated_resources"):
        doc.add_heading("Nguồn học chất lượng", level=2)
        headers = ["Nguồn", "Loại", "Thời gian"]
        rows = []
        for res in data["curated_resources"]:
            if isinstance(res, dict):
                rows.append([
                    res.get("name", ""),
                    res.get("type", ""),
                    f"{res.get('hours', '?')}h"
                ])
            else:
                rows.append([str(res), "", ""])
        if rows:
            add_styled_table(doc, headers, rows)

    # Knowledge map
    if data.get("knowledge_map"):
        doc.add_heading("Knowledge Map", level=2)
        doc.add_paragraph(str(data["knowledge_map"]))

    doc.add_page_break()


def add_section_cong_cu(doc, data):
    """Phần 3: Công Cụ Đúng (20%)."""
    doc.add_heading("Phần 3: 🔧 Công Cụ Đúng (20%) — Luyện Kỹ Năng 20H", level=1)

    # Selected tools
    if data.get("selected_tools"):
        doc.add_heading("Công cụ đã chọn", level=2)
        for tool in data["selected_tools"]:
            if isinstance(tool, dict):
                p = doc.add_paragraph()
                run = p.add_run(f"✅ {tool.get('name', '')}")
                run.font.bold = True
                if tool.get("reason"):
                    p.add_run(f" — {tool['reason']}")
            else:
                doc.add_paragraph(f"✅ {tool}", style='List Bullet')

    # Practice plan
    if data.get("practice_plan"):
        doc.add_heading("Kế hoạch luyện tập", level=2)
        headers = ["Giai đoạn", "Thời gian", "Nội dung"]
        rows = []
        for phase in data["practice_plan"]:
            if isinstance(phase, dict):
                rows.append([
                    phase.get("phase", ""),
                    phase.get("hours", ""),
                    phase.get("content", "")
                ])
            else:
                rows.append([str(phase), "", ""])
        if rows:
            add_styled_table(doc, headers, rows)

    # Milestones
    if data.get("milestones"):
        doc.add_heading("Milestones", level=2)
        for ms in data["milestones"]:
            if isinstance(ms, dict):
                status = "✅" if ms.get("status") == "done" else "⬜"
                doc.add_paragraph(
                    f"{status} {ms.get('name', '')} — Deadline: {ms.get('deadline', '')}",
                    style='List Bullet'
                )
            else:
                doc.add_paragraph(str(ms), style='List Bullet')

    # Proficiency
    if data.get("proficiency"):
        doc.add_heading("Proficiency Level", level=2)
        p = doc.add_paragraph()
        run = p.add_run(f"Mức độ hiện tại: {data['proficiency']}")
        run.font.bold = True
        run.font.size = Pt(14)

    doc.add_page_break()


def add_section_kien_truc(doc, data):
    """Phần Kiến Trúc Kết Hợp: Laptop + Router USB."""
    doc.add_heading("Phần 4: 🏗️ Kiến Trúc Kết Hợp — Demo Đồ Án", level=1)

    # Title & description
    doc.add_heading(data.get("tieu_de", "Kiến Trúc Kết Hợp"), level=2)
    if data.get("mo_ta"):
        p = doc.add_paragraph()
        run = p.add_run(data["mo_ta"])
        run.font.italic = True

    # Architecture diagram (text-based)
    doc.add_heading("Sơ đồ kiến trúc", level=2)
    diagram = (
        "┌──────────────────────────────────┐\n"
        "│  WiFi Router (có USB + ổ cứng)   │ ← Phát WiFi + Lưu trữ file\n"
        "│  IP: 192.168.1.1                 │\n"
        "└──────────────┬───────────────────┘\n"
        "               │ WiFi\n"
        "┌──────────────┴───────────────────┐\n"
        "│  LAPTOP (chạy FastAPI Server)    │ ← Web App + Logic + Database\n"
        "│  IP: 192.168.1.100               │\n"
        "│  Mount USB Router qua Samba      │\n"
        "└──────────────┬───────────────────┘\n"
        "               │ WiFi\n"
        "        ┌──────┴──────┐\n"
        "   📱 Phone       💻 Laptop khác\n"
        "   → 192.168.1.100   → 192.168.1.100"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Components
    if data.get("thanh_phan"):
        doc.add_heading("Thành phần hệ thống", level=2)
        for item in data["thanh_phan"]:
            doc.add_paragraph(str(item), style='List Bullet')

    # Data flow
    if data.get("luong_du_lieu"):
        doc.add_heading("Luồng dữ liệu", level=2)
        p = doc.add_paragraph()
        run = p.add_run(data["luong_du_lieu"])
        run.font.bold = True

    # Demo steps
    if data.get("demo_do_an"):
        doc.add_heading("Các bước Demo Đồ Án", level=2)
        for i, step in enumerate(data["demo_do_an"], 1):
            doc.add_paragraph(str(step), style='List Number')

    # Cost comparison
    if data.get("chi_phi"):
        doc.add_heading("So sánh chi phí", level=2)
        chi_phi = data["chi_phi"]
        rows = [
            ["🎓 Demo Đồ Án", chi_phi.get("demo_do_an", "")],
            ["🏢 Production (nâng cấp sau)", chi_phi.get("production_nang_cap", "")],
        ]
        add_styled_table(doc, ["Kịch bản", "Chi phí"], rows)

    # Advantages
    if data.get("uu_diem"):
        doc.add_heading("Ưu điểm kiến trúc kết hợp", level=2)
        for item in data["uu_diem"]:
            doc.add_paragraph(f"✅ {item}", style='List Bullet')

    doc.add_page_break()


def add_section_summary(doc, topic, data):
    """Phần Tổng kết."""
    doc.add_heading("Phần 5: Tổng Kết & Kế Hoạch Tiếp Theo", level=1)

    doc.add_heading("Tóm tắt", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Chủ đề phân tích: ").font.bold = True
    p.add_run(topic)

    # Summary stats
    doc.add_paragraph()
    doc.add_heading("Thống kê", level=2)

    tu_duy = data.get("tu_duy_dung", {})
    kien_thuc = data.get("kien_thuc_dung", {})
    cong_cu = data.get("cong_cu_dung", {})

    stats = [
        ["🧠 Tư Duy Đúng", f"{len(tu_duy.get('nguyen_ly', []))} nguyên lý, "
                           f"{len(tu_duy.get('mental_models', []))} mental models"],
        ["📘 Kiến Thức Đúng", f"{len(kien_thuc.get('key_concepts', []))} concepts, "
                              f"{len(kien_thuc.get('curated_resources', []))} nguồn học"],
        ["🔧 Công Cụ Đúng", f"{len(cong_cu.get('selected_tools', []))} công cụ, "
                            f"{len(cong_cu.get('milestones', []))} milestones"],
    ]
    add_styled_table(doc, ["Yếu tố", "Kết quả"], stats)

    # Next steps
    doc.add_heading("Kế hoạch tiếp theo (Công Thức Tiến Bộ)", level=2)
    steps = [
        ("📚 Level 1: HỌC + HÀNH = 50%", "Tạo learning plan và thực hành hands-on"),
        ("🎓 Level 2: DẠY LẠI = 90%", "Tạo nội dung chia sẻ (blog, tutorial, video)"),
        ("💰 Level 3: KIẾM TIỀN = 100%", "Monetize kỹ năng (freelance, product, course)"),
    ]
    for title, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        p.add_run(f"\n{desc}")


def generate_docx(data, output_path):
    """Generate DOCX document from collected knowledge."""
    doc = Document()
    setup_styles(doc)

    topic = data.get("topic", "Chưa xác định")
    date_str = data.get("date", datetime.now().strftime("%d/%m/%Y"))

    # Build document
    add_cover_page(doc, topic, date_str)
    add_toc_placeholder(doc)

    if data.get("tu_duy_dung"):
        add_section_tu_duy(doc, data["tu_duy_dung"])

    if data.get("kien_thuc_dung"):
        add_section_kien_thuc(doc, data["kien_thuc_dung"])

    if data.get("cong_cu_dung"):
        add_section_cong_cu(doc, data["cong_cu_dung"])

    if data.get("kien_truc_ket_hop"):
        add_section_kien_truc(doc, data["kien_truc_ket_hop"])

    add_section_summary(doc, topic, data)

    # Save
    doc.save(output_path)
    print(f"✅ DOCX đã được tạo: {output_path}")
    print(f"   Chủ đề: {topic}")
    print(f"   Ngày: {date_str}")


def get_sample_data():
    """Dữ liệu mẫu để test."""
    return {
        "topic": "Trở thành Full-Stack Developer trong 6 tháng",
        "date": datetime.now().strftime("%d/%m/%Y"),
        "tu_duy_dung": {
            "nguyen_ly": [
                "Web development = Client + Server + Database",
                "HTTP là giao thức nền tảng cho mọi web app",
                "JavaScript chạy được cả frontend và backend (Node.js)",
                "Mọi framework đều xây trên HTML/CSS/JS cơ bản"
            ],
            "mental_models": [
                "Pareto: 20% kiến thức JS core = 80% ứng dụng",
                "Inversion: Tránh framework hopping, thiếu fundamentals",
                "Circle of Competence: Focus frontend trước → backend sau"
            ],
            "first_principles": [
                "Web app = Input → Processing → Output → Display",
                "User chỉ quan tâm: nhanh, đẹp, hoạt động đúng",
                "Security và performance là non-negotiable"
            ],
            "five_whys": [
                "Muốn làm full-stack developer",
                "Vì muốn tự xây sản phẩm từ A-Z",
                "Vì không muốn phụ thuộc vào người khác",
                "Vì muốn kiếm tiền từ sản phẩm của mình",
                "GỐC RỄ: Muốn tự do tài chính và sáng tạo"
            ],
            "pareto_top20": [
                "JavaScript/TypeScript (dùng cho cả FE và BE)",
                "React (frontend phổ biến nhất)",
                "Node.js + Express (backend nhanh nhất để bắt đầu)",
                "PostgreSQL (database tin cậy)"
            ],
            "removed_assumptions": [
                "Phải học nhiều ngôn ngữ khác nhau",
                "Cần bằng cấp mới được nhận làm developer",
                "Framework mới = tốt hơn framework cũ"
            ]
        },
        "kien_thuc_dung": {
            "knowledge_gaps": [
                {"field": "JavaScript", "current": 2, "target": 4, "gap": 2},
                {"field": "React", "current": 1, "target": 4, "gap": 3},
                {"field": "Node.js", "current": 1, "target": 3, "gap": 2},
                {"field": "Database", "current": 2, "target": 3, "gap": 1}
            ],
            "curated_resources": [
                {"name": "The Odin Project", "type": "curriculum", "hours": 100},
                {"name": "JavaScript.info", "type": "documentation", "hours": 40},
                {"name": "React Official Docs", "type": "documentation", "hours": 20},
                {"name": "Fullstack Open (Helsinki)", "type": "course", "hours": 60}
            ],
            "key_concepts": [
                {"name": "DOM Manipulation", "level": "core", "description": "Giao tiếp JS với HTML"},
                {"name": "Async/Await", "level": "core", "description": "Xử lý bất đồng bộ"},
                {"name": "REST API", "level": "core", "description": "Giao tiếp client-server"},
                {"name": "State Management", "level": "supporting", "description": "Quản lý state trong React"},
                {"name": "Authentication", "level": "supporting", "description": "Xác thực người dùng"}
            ],
            "knowledge_map": "JS Fundamentals → DOM → React → Node.js → Database → Full-Stack App"
        },
        "cong_cu_dung": {
            "selected_tools": [
                {"name": "VS Code", "score": 24, "reason": "IDE phổ biến nhất, ecosystem mạnh"},
                {"name": "Git + GitHub", "score": 22, "reason": "Version control, portfolio"},
                {"name": "Docker", "score": 18, "reason": "Containerization cho deployment"}
            ],
            "practice_plan": [
                {"phase": "Setup", "hours": "1-2h", "content": "Cài VS Code, Node.js, Git"},
                {"phase": "JS Fundamentals", "hours": "3-6h", "content": "Variables, functions, DOM"},
                {"phase": "React Basics", "hours": "7-12h", "content": "Components, state, effects"},
                {"phase": "Full-Stack Project", "hours": "13-17h", "content": "Todo app with API"},
                {"phase": "Deployment", "hours": "18-20h", "content": "Deploy lên Vercel + Railway"}
            ],
            "milestones": [
                {"id": "M1", "name": "Dev environment ready", "deadline": "Hour 2", "status": "done"},
                {"id": "M2", "name": "JS fundamentals solid", "deadline": "Hour 6", "status": "done"},
                {"id": "M3", "name": "React mini-app working", "deadline": "Hour 12", "status": "in_progress"},
                {"id": "M4", "name": "Full-stack app deployed", "deadline": "Hour 17", "status": "pending"},
                {"id": "M5", "name": "Portfolio complete", "deadline": "Hour 20", "status": "pending"}
            ],
            "proficiency": "Intermediate (đang tiến tới Proficient)"
        }
    }


def main():
    parser = argparse.ArgumentParser(
        description="Xuất DOCX từ kiến thức Tam Giác Thành Công"
    )
    parser.add_argument("--input", "-i", help="Path to JSON input file")
    parser.add_argument("--output", "-o", default="output.docx", help="Output DOCX path")
    parser.add_argument("--sample", action="store_true", help="Generate sample DOCX with test data")

    args = parser.parse_args()

    if args.sample:
        data = get_sample_data()
        output = args.output if args.output != "output.docx" else "sample_tam_giac.docx"
        generate_docx(data, output)
    elif args.input:
        if not os.path.exists(args.input):
            print(f"❌ File không tồn tại: {args.input}")
            sys.exit(1)
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
        generate_docx(data, args.output)
    else:
        parser.print_help()
        print("\n💡 Tip: Dùng --sample để tạo DOCX mẫu.")


if __name__ == "__main__":
    main()
